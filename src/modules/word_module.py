import os
import re
import docx
import zipfile
from src.core.base_module import BaseDocumentModule
from src.core.registry import ModuleRegistry

class WordModule(BaseDocumentModule):
    @property
    def name(self) -> str:
        return "Word"

    @property
    def file_extensions(self) -> list[str]:
        return [".docx"]

    @property
    def required_dependencies(self) -> list[str]:
        return ["python-docx", "markitdown"]

    def load_to_markdown(self, file_path: str) -> str:
        """Extracts Word .docx to clean Markdown text, preserving tables, headings, bold/italic styles, and lists."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            from docx.oxml.text.paragraph import CT_P
            from docx.oxml.table import CT_Tbl
            from docx.text.paragraph import Paragraph
            from docx.table import Table

            if not zipfile.is_zipfile(file_path):
                return "⚠️ File Validation Error: Invalid file. Please select a valid DOCX document."

            doc = docx.Document(file_path)
            parts = []

            def iter_block_items(parent):
                parent_elm = parent.element.body
                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            from docx.oxml.ns import qn
            from src.core.converters import wrap_text_style
            from src.services.media_asset_manager import MediaAssetManager
            asset_mgr = MediaAssetManager()

            def extract_and_get_link(run) -> str:
                xml = run._element.xml
                rId_match = re.search(r'(?:r:embed|embed)="([^"]+)"', xml)
                if rId_match:
                    rId = rId_match.group(1)
                    try:
                        image_part = doc.part.related_parts[rId]
                        if hasattr(image_part, "image"):
                            image_bytes = image_part.image.blob
                            ext = "png"
                            if hasattr(image_part.image, "ext") and image_part.image.ext:
                                ext = image_part.image.ext
                            elif hasattr(image_part.image, "content_type") and image_part.image.content_type:
                                ct = image_part.image.content_type
                                if "/" in ct:
                                    ext = ct.split("/")[1]
                            
                            filename = f"image_{rId}.{ext}"
                            virtual_uri = asset_mgr.register_image(image_bytes, filename)
                            return f"![image]({virtual_uri})"
                    except Exception as e:
                        print(f"[DEBUG] Failed to extract run image with rId {rId}: {e}")
                return "[image]"

            for block in iter_block_items(doc):
                if isinstance(block, Paragraph):
                    style_name = (block.style.name or "").lower() if block.style else ""
                    
                    def run_contains_image(run):
                        xml = getattr(run._element, "xml", "")
                        return "<w:drawing" in xml or "<w:pict" in xml or "<a:blip" in xml

                    para_parts = []
                    for child in block._element:
                        if child.tag == qn('w:r'):
                            run = docx.text.run.Run(child, block)
                            if run_contains_image(run):
                                para_parts.append(extract_and_get_link(run))
                                continue
                            if not run.text:
                                continue
                            formatted = wrap_text_style(
                                run.text,
                                bold=run.bold,
                                italic=run.italic,
                                strike=run.font.strike,
                                underline=run.font.underline
                            )
                            para_parts.append(formatted)
                        elif child.tag == qn('w:hyperlink'):
                            r_id = child.get(qn('r:id'))
                            url = ""
                            if r_id:
                                try:
                                    url = block.part.rels[r_id].target_ref
                                except Exception:
                                    pass
                            link_parts = []
                            for sub_child in child:
                                  if sub_child.tag == qn('w:r'):
                                     run = docx.text.run.Run(sub_child, block)
                                     if run_contains_image(run):
                                         link_parts.append(extract_and_get_link(run))
                                         continue
                                     if not run.text:
                                         continue
                                     formatted = wrap_text_style(
                                         run.text,
                                         bold=run.bold,
                                         italic=run.italic,
                                         strike=run.font.strike,
                                         underline=run.font.underline
                                     )
                                     link_parts.append(formatted)
                            link_text = "".join(link_parts)
                            if link_text:
                                if url:
                                    para_parts.append(f"[{link_text}]({url})")
                                else:
                                    para_parts.append(link_text)
                    
                    para_text = "".join(para_parts).strip()
                    if not para_text:
                        if block.text.strip():
                            para_text = block.text.strip()
                        else:
                            parts.append("")
                            continue
                    
                    is_heading = style_name.startswith("heading") or re.match(r"^(đề mục|tiêu đề)\s*\d", style_name)

                    if is_heading:
                        m = re.search(r"\d+", style_name)
                        level = int(m.group(0)) if m else 1
                        parts.append("#" * level + " " + para_text)
                    elif "bullet" in style_name:
                        m = re.search(r"\d+", style_name)
                        level = int(m.group(0)) if m else 1
                        indent = "  " * (level - 1)
                        parts.append(indent + "- " + para_text)
                    elif "number" in style_name or style_name.startswith("list"):
                        m = re.search(r"\d+", style_name)
                        level = int(m.group(0)) if m else 1
                        indent = "  " * (level - 1)
                        parts.append(indent + "1. " + para_text)
                    else:
                        parts.append(para_text)

                elif isinstance(block, Table):
                    table_parts = []
                    rows_data = []
                    for row in block.rows:
                        row_cells = []
                        for cell in row.cells:
                            cell_text = cell.text.strip().replace("\n", " ").replace("|", "\\|")
                            row_cells.append(cell_text)
                        rows_data.append(row_cells)
                    
                    if rows_data:
                        header = "| " + " | ".join(rows_data[0]) + " |"
                        sep = "| " + " | ".join("---" for _ in rows_data[0]) + " |"
                        table_parts.append(header)
                        table_parts.append(sep)
                        for row in rows_data[1:]:
                            table_parts.append("| " + " | ".join(row) + " |")
                        parts.append("\n".join(table_parts))

            return "\n\n".join(parts)
        except Exception as e:
            import sys
            print(f"[DEBUG] Custom Word parsing failed: {e}. Falling back to markitdown.", file=sys.stderr)
            try:
                from markitdown import MarkItDown
                md = MarkItDown()
                result = md.convert(file_path)
                if not result or not result.text_content:
                    return "*(Empty Word Document)*"
                return result.text_content
            except Exception as inner_e:
                raise RuntimeError(f"Word Ingestion Error: Failed to extract text layer from DOCX file. Detail: {str(inner_e)}")

    def save_from_markdown(self, markdown_content: str, out_path: str) -> str:
        """Converts Markdown text to formatted Word document."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        FONT = "Arial"
        HEADING_SIZES = {1: 20, 2: 16, 3: 13, 4: 12, 5: 11, 6: 11}
        HEADING_COLORS = {i: "404040" for i in range(1, 7)}

        def set_font(run, size=11, bold=False, color=None):
            run.font.name = FONT
            run.font.size = Pt(size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = RGBColor.from_string(color)

        def add_hyperlink_run(paragraph, text, url, size=11, bold=False, italic=False, strike=False, underline=True, color="0000FF"):
            part = paragraph.part
            r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)

            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')

            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), FONT)
            rFonts.set(qn('w:hAnsi'), FONT)
            rPr.append(rFonts)

            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(size * 2)))
            rPr.append(sz)

            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)
            if italic:
                i = OxmlElement('w:i')
                rPr.append(i)
            if strike:
                strike_el = OxmlElement('w:strike')
                rPr.append(strike_el)
            if underline:
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)

            if color:
                c = OxmlElement('w:color')
                c.set(qn('w:val'), color)
                rPr.append(c)

            new_run.append(rPr)
            
            text_node = OxmlElement('w:t')
            text_node.text = text
            new_run.append(text_node)

            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
            return hyperlink

        def add_formatted_runs(paragraph, text, size=11, default_bold=False, default_color=None):
            from src.core.converters import parse_inline
            segments = parse_inline(text, bold=default_bold)
            for seg in segments:
                if seg.url:
                    add_hyperlink_run(
                        paragraph, 
                        seg.text, 
                        seg.url, 
                        size=size, 
                        bold=seg.bold, 
                        italic=seg.italic, 
                        strike=seg.strike, 
                        underline=True, 
                        color="0000FF"
                    )
                else:
                    run = paragraph.add_run(seg.text)
                    run.font.name = "Consolas" if seg.code else FONT
                    run.font.size = Pt(size - 1) if seg.code else Pt(size)
                    run.font.bold = seg.bold
                    run.font.italic = seg.italic
                    run.font.strike = seg.strike
                    run.font.underline = seg.underline
                    if seg.code:
                        run.font.color.rgb = RGBColor(0xA5, 0x2A, 0x2A)
                    elif default_color:
                        run.font.color.rgb = RGBColor.from_string(default_color)

        def add_paragraph_with_font(doc, text, size=11, bold=False, color=None, style=None):
            p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            add_formatted_runs(p, text, size=size, default_bold=bold, default_color=color)
            return p

        lines = markdown_content.splitlines()
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = FONT # type: ignore
        style.font.size = Pt(11) # type: ignore

        i = 0
        while i < len(lines):
            line = lines[i].rstrip("\r\n")

            m = re.match(r"^(#{1,6})\s+(.*)", line)
            if m:
                level = len(m.group(1))
                heading = doc.add_heading('', level=min(level, 9))
                add_formatted_runs(heading, m.group(2), size=HEADING_SIZES[level], default_bold=True, default_color=HEADING_COLORS[level])
                heading.paragraph_format.space_before = Pt(10)
                heading.paragraph_format.space_after = Pt(4)
                i += 1
                continue

            if "|" in line and not re.match(r"^[\|\s\-:]+$", line.strip()):
                table_lines = []
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i].strip())
                    i += 1
                data_lines = [l for l in table_lines if not re.match(r"^[\|\s\-:]+$", l)]
                if len(data_lines) >= 2:
                    rows = [[c.strip() for c in l.split("|") if c.strip()] for l in data_lines]
                    max_cols = max(len(r) for r in rows)
                    rows = [r + [""] * (max_cols - len(r)) for r in rows]
                    tbl = doc.add_table(rows=len(rows), cols=max_cols)
                    tbl.style = "Table Grid"
                    for r_idx, row_data in enumerate(rows):
                        for c_idx, cell_text in enumerate(row_data):
                            cell = tbl.cell(r_idx, c_idx)
                            cell.text = ""
                            run = cell.paragraphs[0].add_run(cell_text)
                            if r_idx == 0:
                                set_font(run, size=11, bold=True, color="FFFFFF")
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement("w:shd")
                                shd.set(qn("w:fill"), "4472C4")
                                shd.set(qn("w:color"), "auto")
                                shd.set(qn("w:val"), "clear")
                                tcPr.append(shd)
                            else:
                                set_font(run, size=11)
                    doc.add_paragraph()
                continue

            m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
            if m:
                indent_spaces = len(m.group(1))
                level = (indent_spaces // 2) + 1
                style_name = f"List Bullet {level}" if level > 1 else "List Bullet"
                if style_name not in doc.styles:
                    style_name = "List Bullet"
                add_paragraph_with_font(doc, m.group(2), style=style_name)
                i += 1
                continue

            m = re.match(r"^(\s*)\d+\.\s+(.*)", line)
            if m:
                indent_spaces = len(m.group(1))
                level = (indent_spaces // 2) + 1
                style_name = f"List Number {level}" if level > 1 else "List Number"
                if style_name not in doc.styles:
                    style_name = "List Number"
                add_paragraph_with_font(doc, m.group(2), style=style_name)
                i += 1
                continue

            if re.match(r"^[-*_]{3,}$", line.strip()):
                p = doc.add_paragraph("─" * 50)
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                i += 1
                continue

            text = re.sub(r"`(.*?)`", r"\1", line)
            if text.strip():
                add_paragraph_with_font(doc, text)
            i += 1

        doc.save(out_path)
        return f"Word document created successfully -> {os.path.basename(out_path)}"

ModuleRegistry.register(WordModule())
