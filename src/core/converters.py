import os
import re
import pandas as pd

def parse_md_tables(content: str) -> list[tuple[str, pd.DataFrame]]:
    tables, lines, i = [], content.split("\n"), 0
    while i < len(lines):
        line = lines[i].strip()
        if "|" in line and not re.match(r"^[\|\s\-:]+$", line):
            table_name = f"Sheet{len(tables)+1}"
            for j in range(i-1, max(i-5, -1), -1):
                prev = lines[j].strip()
                if prev.startswith("#"):
                    table_name = re.sub(r"^#+\s*", "", prev)[:31]
                    break
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            data_lines = [l for l in table_lines if not re.match(r"^[\|\s\-:]+$", l)]
            if len(data_lines) < 2:
                continue
            rows = [[c.strip() for c in l.split("|") if c.strip()] for l in data_lines]
            max_cols = max(len(r) for r in rows)
            rows = [r + [""] * (max_cols - len(r)) for r in rows]
            df = pd.DataFrame(rows[1:], columns=rows[0])
            tables.append((table_name, df))
        else:
            i += 1
    return tables


def md_to_excel_from_text(content: str, out_path: str) -> str:
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    tables = parse_md_tables(content)
    if not tables:
        return "No tables found in the Markdown content."

    seen = {}
    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        for name, df in tables:
            key = name
            if key in seen:
                seen[key] += 1
                key = f"{name}_{seen[key]}"
            else:
                seen[name] = 0
            df.to_excel(writer, sheet_name=key, index=False)

            ws = writer.sheets[key]

            header_fill = PatternFill("solid", fgColor="4472C4")
            header_font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
            body_font = Font(name="Arial", size=11)
            thin = Side(border_style="thin", color="D9D9D9")
            thin_border  = Border(left=thin, right=thin, top=thin, bottom=thin)

            for row_idx, row in enumerate(ws.iter_rows(), start=1):
                for cell in row:
                    cell.border = thin_border
                    if row_idx == 1:
                        cell.fill = header_fill
                        cell.font = header_font
                        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                    else:
                        cell.font = body_font
                        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

            for col in ws.columns:
                max_len = 0
                for cell in col:
                    if cell.value:
                        max_len = max(max_len, len(str(cell.value)))
                ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 5, 50)

            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions

    return f"Exported {len(tables)} sheet(s) successfully -> {os.path.basename(out_path)}"


def md_to_word_from_text(content: str, out_path: str) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FONT = "Arial"
    HEADING_SIZES = {1: 20, 2: 16, 3: 13, 4: 12, 5: 11, 6: 11}
    HEADING_COLORS = {i: "404040" for i in range(1, 7)}

    def set_font(run, size=11, bold=False, color=None):
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    def add_paragraph_with_font(doc, text, size=11, bold=False, color=None, style=None):
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                set_font(run, size=size, bold=True, color=color)
            else:
                run = p.add_run(part)
                set_font(run, size=size, bold=bold, color=color)
        return p

    lines = content.splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT # type: ignore
    style.font.size = Pt(11) # type: ignore

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\r\n")

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(m.group(2), level=min(level, 9))
            for run in heading.runs:
                set_font(run, size=HEADING_SIZES[level], bold=True, color=HEADING_COLORS[level])
            heading.paragraph_format.space_before = Pt(10)
            heading.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Table
        if "|" in line and not re.match(r"^[\|\s\-:]+$", line.strip()):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            data_lines = [l for l in table_lines if not re.match(r"^[\|\s\-:]+$", l)]
            if len(data_lines) >= 2:
                rows = [[c.strip() for c in l.split("|") if c.strip()] for l in data_lines]
                max_cols = max(len(r) for r in rows)
                rows = [r + [""] * (max_cols - len(r)) for r in rows]
                tbl = doc.add_table(rows=len(rows), cols=max_cols)
                tbl.style = "Table Grid"
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_data):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(cell_text)
                        if r_idx == 0:
                            set_font(run, size=11, bold=True, color="FFFFFF")
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            shd = OxmlElement("w:shd")
                            shd.set(qn("w:fill"), "4472C4")
                            shd.set(qn("w:color"), "auto")
                            shd.set(qn("w:val"), "clear")
                            tcPr.append(shd)
                        else:
                            set_font(run, size=11)
                doc.add_paragraph()  # spacing
            continue

        # Unordered list
        m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if m:
            add_paragraph_with_font(doc, m.group(2), style="List Bullet")
            i += 1
            continue

        # Ordered list
        m = re.match(r"^\s*\d+\.\s+(.*)", line)
        if m:
            add_paragraph_with_font(doc, m.group(1), style="List Number")
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", line.strip()):
            p = doc.add_paragraph("─" * 50)
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            i += 1
            continue

        # Normal paragraph
        text = re.sub(r"`(.*?)`", r"\1", line)
        if text.strip():
            add_paragraph_with_font(doc, text)
        i += 1

    doc.save(out_path)
    return f"Word document created successfully -> {os.path.basename(out_path)}"


def save_markdown_from_text(content: str, out_path: str) -> str:
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(content)
    return f"Markdown file saved successfully -> {os.path.basename(out_path)}"
