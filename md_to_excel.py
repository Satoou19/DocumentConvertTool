"""
Document Converter Workspace - Premium Version
Supports: MD→Excel, MD→Word, Excel→MD, Word→MD
Features: Dual-Pane Live Editor, Overview, Auto-Extraction on Load, Drag & Drop
"""

import sys
import re
import os
import threading
import pandas as pd
import customtkinter as ctk
from tkinter import filedialog

# Prevent UnicodeEncodeError when printing emojis/UTF-8 characters to standard output in Windows console
if sys.stdout is not None:
    try:
        sys.stdout.reconfigure(encoding='utf-8', errors='backslashreplace')
    except AttributeError:
        pass
if sys.stderr is not None:
    try:
        sys.stderr.reconfigure(encoding='utf-8', errors='backslashreplace')
    except AttributeError:
        pass

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
except ImportError:
    HAS_DND = False
    print("[WARN] tkinterdnd2 not found, drag & drop disabled")


# ── Conversion modes config ───────────────────────────────────────────────────

MODES = {
    "MD → Excel":  {"in_ext": ".md",   "out_ext": ".xlsx", "in_label": "File .md",   "out_label": "Lưu .xlsx"},
    "MD → Word":   {"in_ext": ".md",   "out_ext": ".docx", "in_label": "File .md",   "out_label": "Lưu .docx"},
    "Excel → MD":  {"in_ext": ".xlsx", "out_ext": ".md",   "in_label": "File .xlsx", "out_label": "Lưu .md"},
    "Word → MD":   {"in_ext": ".docx", "out_ext": ".md",   "in_label": "File .docx", "out_label": "Lưu .md"},
}

IN_FILETYPES = {
    ".md":   [("Markdown", "*.md"), ("All Files", "*.*")],
    ".xlsx": [("Excel", "*.xlsx *.xls"), ("All Files", "*.*")],
    ".docx": [("Word", "*.docx"), ("All Files", "*.*")],
}

OUT_FILETYPES = {
    ".xlsx": [("Excel", "*.xlsx")],
    ".docx": [("Word",  "*.docx")],
    ".md":   [("Markdown", "*.md")],
}


# ── Extractor Helpers (File to Markdown text) ───────────────────────────────

def extract_excel_to_md(in_path: str) -> str:
    """Extracts Excel sheets into clean Markdown tables."""
    xl = pd.ExcelFile(in_path)
    parts = []
    for sheet in xl.sheet_names:
        df = xl.parse(sheet).fillna("")
        parts.append(f"## {sheet}\n")
        # Generate Markdown Table representation
        if not df.empty:
            header = "| " + " | ".join(str(c) for c in df.columns) + " |"
            sep    = "| " + " | ".join("---" for _ in df.columns) + " |"
            parts.append(header)
            parts.append(sep)
            for _, row in df.iterrows():
                parts.append("| " + " | ".join(str(v) for v in row) + " |")
        else:
            parts.append("*(Bảng trống)*")
        parts.append("")
    return "\n".join(parts)


def extract_word_to_md(in_path: str) -> str:
    """Extracts Word .docx to clean Markdown text using Mammoth."""
    import mammoth
    with open(in_path, "rb") as f:
        result = mammoth.convert_to_markdown(f)
    markdown = result.value
    # Clean up double backslashes Mammoth often introduces
    markdown = re.sub(r"\\([.\-_~*`\[\]()#+!{}])", r"\1", markdown)
    
    # Clean list numbers
    lines = markdown.splitlines()
    global_counter = 1
    final_lines = []
    for line in lines:
        stripped = line.strip()
        if re.match(r"^1\.\s+", stripped):
            replaced_line = re.sub(
                r"^(.*?)1\.\s+", 
                lambda m: f"{m.group(1)}{global_counter}. ", 
                line, 
                count=1
            )
            final_lines.append(replaced_line)
            global_counter += 1
        else:
            final_lines.append(line)
    return "\n".join(final_lines)


# ── Parsers & Converters (Markdown text to Target files) ──────────────────────

def parse_md_tables(content: str) -> list[tuple[str, pd.DataFrame]]:
    tables, lines, i = [], content.split("\n"), 0
    while i < len(lines):
        line = lines[i].strip()
        if "|" in line and not re.match(r"^[\|\s\-:]+$", line):
            table_name = f"Sheet{len(tables)+1}"
            for j in range(i-1, max(i-5, -1), -1):
                prev = lines[j].strip()
                if prev.startswith("#"):
                    table_name = re.sub(r"^#+\s*", "", prev)[:31]
                    break
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            data_lines = [l for l in table_lines if not re.match(r"^[\|\s\-:]+$", l)]
            if len(data_lines) < 2:
                continue
            rows = [[c.strip() for c in l.split("|") if c.strip()] for l in data_lines]
            max_cols = max(len(r) for r in rows)
            rows = [r + [""] * (max_cols - len(r)) for r in rows]
            df = pd.DataFrame(rows[1:], columns=rows[0])
            tables.append((table_name, df))
        else:
            i += 1
    return tables


def md_to_excel_from_text(content: str, out_path: str) -> str:
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    tables = parse_md_tables(content)
    if not tables:
        return "❌ Không tìm thấy bảng biểu (table) nào trong trình soạn thảo Markdown."

    seen = {}
    with pd.ExcelWriter(out_path, engine="openpyxl") as writer:
        for name, df in tables:
            key = name
            if key in seen:
                seen[key] += 1
                key = f"{name}_{seen[key]}"
            else:
                seen[name] = 0
            df.to_excel(writer, sheet_name=key, index=False)

            ws = writer.sheets[key]

            header_fill = PatternFill("solid", fgColor="4472C4")
            header_font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
            body_font = Font(name="Arial", size=11)
            thin = Side(border_style="thin", color="D9D9D9")
            thin_border  = Border(left=thin, right=thin, top=thin, bottom=thin)

            for row_idx, row in enumerate(ws.iter_rows(), start=1):
                for cell in row:
                    cell.border = thin_border
                    if row_idx == 1:
                          cell.fill = header_fill
                          cell.font = header_font
                          cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                    else:
                          cell.font = body_font
                          cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)

            for col in ws.columns:
                max_len = 0
                for cell in col:
                    if cell.value:
                        max_len = max(max_len, len(str(cell.value)))
                ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 5, 50)

            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions

    return f"✅ Xuất {len(tables)} sheet(s) thành công → {os.path.basename(out_path)}"


def md_to_word_from_text(content: str, out_path: str) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    FONT = "Arial"
    HEADING_SIZES = {1: 20, 2: 16, 3: 13, 4: 12, 5: 11, 6: 11}
    HEADING_COLORS = {i: "404040" for i in range(1, 7)}

    def set_font(run, size=11, bold=False, color=None):
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    def add_paragraph_with_font(doc, text, size=11, bold=False, color=None, style=None):
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                set_font(run, size=size, bold=True, color=color)
            else:
                run = p.add_run(part)
                set_font(run, size=size, bold=bold, color=color)
        return p

    lines = content.splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT # type: ignore
    style.font.size = Pt(11) # type: ignore

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\r\n")

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(m.group(2), level=min(level, 9))
            for run in heading.runs:
                set_font(run, size=HEADING_SIZES[level], bold=True, color=HEADING_COLORS[level])
            heading.paragraph_format.space_before = Pt(10)
            heading.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Table
        if "|" in line and not re.match(r"^[\|\s\-:]+$", line.strip()):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            data_lines = [l for l in table_lines if not re.match(r"^[\|\s\-:]+$", l)]
            if len(data_lines) >= 2:
                rows = [[c.strip() for c in l.split("|") if c.strip()] for l in data_lines]
                max_cols = max(len(r) for r in rows)
                rows = [r + [""] * (max_cols - len(r)) for r in rows]
                tbl = doc.add_table(rows=len(rows), cols=max_cols)
                tbl.style = "Table Grid"
                for r_idx, row_data in enumerate(rows):
                    for c_idx, cell_text in enumerate(row_data):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = ""
                        run = cell.paragraphs[0].add_run(cell_text)
                        if r_idx == 0:
                            set_font(run, size=11, bold=True, color="FFFFFF")
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            shd = OxmlElement("w:shd")
                            shd.set(qn("w:fill"), "4472C4")
                            shd.set(qn("w:color"), "auto")
                            shd.set(qn("w:val"), "clear")
                            tcPr.append(shd)
                        else:
                            set_font(run, size=11)
                doc.add_paragraph()  # spacing
            continue

        # Unordered list
        m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if m:
            add_paragraph_with_font(doc, m.group(2), style="List Bullet")
            i += 1
            continue

        # Ordered list
        m = re.match(r"^\s*\d+\.\s+(.*)", line)
        if m:
            add_paragraph_with_font(doc, m.group(1), style="List Number")
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", line.strip()):
            p = doc.add_paragraph("─" * 50)
            p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            i += 1
            continue

        # Normal paragraph
        text = re.sub(r"`(.*?)`", r"\1", line)
        if text.strip():
            add_paragraph_with_font(doc, text)
        i += 1

    doc.save(out_path)
    return f"✅ Đã tạo Word thành công → {os.path.basename(out_path)}"


def save_markdown_from_text(content: str, out_path: str) -> str:
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(content)
    return f"✅ Đã lưu file Markdown thành công → {os.path.basename(out_path)}"


# ── Premium GUI Workspace ─────────────────────────────────────────────────────

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

BaseClass = TkinterDnD.Tk if HAS_DND else ctk.CTk


class App(BaseClass): # type: ignore
    def __init__(self):
        super().__init__()
        self.title("Document Converter Workspace")
        self.geometry("1200x780")
        self.resizable(True, True)
        
        # Configure app-wide variables
        self.in_path = ctk.StringVar()
        self.out_path = ctk.StringVar()
        self.mode_var = ctk.StringVar(value="MD → Excel")
        
        self._build_ui()
        self._on_mode_change()

    def _build_ui(self):
        # Grid Configuration for main workspace
        self.rowconfigure(0, weight=0) # Header
        self.rowconfigure(1, weight=1) # Main workspace split
        self.columnconfigure(0, weight=1)

        # 1. Sleek Header Panel
        header_frame = ctk.CTkFrame(self, fg_color="#1e1e24", height=70, corner_radius=0)
        header_frame.grid(row=0, column=0, sticky="ew", padx=0, pady=0)
        header_frame.grid_propagate(False)
        
        title_lbl = ctk.CTkLabel(
            header_frame, text="📄 Document Converter Workspace",
            font=ctk.CTkFont(family="Arial", size=20, weight="bold"), text_color="#ffffff"
        )
        title_lbl.pack(side="left", padx=25, pady=10)
        
        subtitle_lbl = ctk.CTkLabel(
            header_frame, text="Workspace soạn thảo & chuyển đổi tài liệu đa năng",
            font=ctk.CTkFont(family="Arial", size=13, slant="italic"), text_color="#8a8a9e"
        )
        subtitle_lbl.pack(side="left", padx=10, pady=16)

        # 2. Main Workspace Split (Left Pane & Right Pane)
        workspace = ctk.CTkFrame(self, fg_color="transparent")
        workspace.grid(row=1, column=0, sticky="nsew", padx=15, pady=15)
        workspace.columnconfigure(0, weight=6) # Left Pane gets more weight
        workspace.columnconfigure(1, weight=5) # Right Pane
        workspace.rowconfigure(0, weight=1)

        # ── LEFT PANE: Input Editor & Overview ────────────────────────────────
        left_pane = ctk.CTkFrame(workspace, fg_color="#18181c", corner_radius=12, border_width=1, border_color="#2c2c35")
        left_pane.grid(row=0, column=0, sticky="nsew", padx=(0, 10), pady=0)
        left_pane.rowconfigure(0, weight=0) # Label
        left_pane.rowconfigure(1, weight=0) # Drag Drop Info
        left_pane.rowconfigure(2, weight=1) # Textbox Editor
        left_pane.rowconfigure(3, weight=0) # Footer stats
        left_pane.columnconfigure(0, weight=1)

        editor_title = ctk.CTkLabel(
            left_pane, text="📝 BẢNG SOẠN THẢO & OVERVIEW ĐẦU VÀO (MARKDOWN / TEXT)",
            font=ctk.CTkFont(family="Arial", size=13, weight="bold"), text_color="#3a86ff"
        )
        editor_title.grid(row=0, column=0, sticky="w", padx=15, pady=(12, 5))

        # File Load Area / Drag Zone
        self.load_bar = ctk.CTkFrame(left_pane, fg_color="#1c1c24", corner_radius=8, height=45)
        self.load_bar.grid(row=1, column=0, sticky="ew", padx=15, pady=5)
        self.load_bar.pack_propagate(False)
        
        self.drop_lbl = ctk.CTkLabel(
            self.load_bar, text="📂 Kéo thả file hoặc nhấn 'Browse' để nạp nội dung...",
            font=ctk.CTkFont(family="Arial", size=12), text_color="#7eb8f5"
        )
        self.drop_lbl.pack(side="left", padx=15, fill="both", expand=True)

        btn_browse_in = ctk.CTkButton(
            self.load_bar, text="Browse", width=75, height=28,
            font=ctk.CTkFont(family="Arial", size=12, weight="bold"), fg_color="#2980b9", hover_color="#3498db",
            command=self._browse_input
        )
        btn_browse_in.pack(side="right", padx=8, pady=8)

        # Register Drag & Drop
        if HAS_DND:
            for w in (self.load_bar, self.drop_lbl):
                w.drop_target_register(DND_FILES) # type: ignore
                w.dnd_bind("<<Drop>>", self._on_drop) # type: ignore

        # The Live Monospace Text Editor
        self.editor = ctk.CTkTextbox(
            left_pane, fg_color="#111115", text_color="#f8f8f2",
            font=ctk.CTkFont(family="Consolas", size=13),
            border_width=1, border_color="#2c2c35", corner_radius=8
        )
        self.editor.grid(row=2, column=0, sticky="nsew", padx=15, pady=8)
        self.editor.bind("<KeyRelease>", self._update_counts)

        # Editor Footer (Stats & Actions)
        editor_footer = ctk.CTkFrame(left_pane, fg_color="transparent")
        editor_footer.grid(row=3, column=0, sticky="ew", padx=15, pady=(5, 12))
        
        self.char_lbl = ctk.CTkLabel(
            editor_footer, text="Ký tự: 0", font=ctk.CTkFont(family="Arial", size=12), text_color="#8a8a9e"
        )
        self.char_lbl.pack(side="left", padx=5)

        self.word_lbl = ctk.CTkLabel(
            editor_footer, text=" |  Từ: 0", font=ctk.CTkFont(family="Arial", size=12), text_color="#8a8a9e"
        )
        self.word_lbl.pack(side="left", padx=5)

        btn_clear = ctk.CTkButton(
            editor_footer, text="Xóa Hết", width=70, height=24, fg_color="#c0392b", hover_color="#e74c3c",
            font=ctk.CTkFont(family="Arial", size=11, weight="bold"),
            command=self._clear_editor
        )
        btn_clear.pack(side="right", padx=5)


        # ── RIGHT PANE: Output Configuration, Preview & Actions ───────────────
        right_pane = ctk.CTkFrame(workspace, fg_color="#18181c", corner_radius=12, border_width=1, border_color="#2c2c35")
        right_pane.grid(row=0, column=1, sticky="nsew", padx=(10, 0), pady=0)
        right_pane.rowconfigure(0, weight=0) # Title
        right_pane.rowconfigure(1, weight=0) # Config panel
        right_pane.rowconfigure(2, weight=1) # Output preview / Logs textbox
        right_pane.rowconfigure(3, weight=0) # Convert Button & Status
        right_pane.columnconfigure(0, weight=1)

        output_title = ctk.CTkLabel(
            right_pane, text="⚡ CẤU HÌNH & XUẤT ĐẦU RA",
            font=ctk.CTkFont(family="Arial", size=13, weight="bold"), text_color="#2ecc71"
        )
        output_title.grid(row=0, column=0, sticky="w", padx=15, pady=(12, 5))

        # Config Panel
        config_frame = ctk.CTkFrame(right_pane, fg_color="#1c1c24", corner_radius=8)
        config_frame.grid(row=1, column=0, sticky="ew", padx=15, pady=5)
        config_frame.columnconfigure(0, weight=1)

        # Mode row
        row_mode = ctk.CTkFrame(config_frame, fg_color="transparent")
        row_mode.pack(fill="x", padx=12, pady=(10, 5))
        ctk.CTkLabel(row_mode, text="Chế độ:", width=70, anchor="w", font=ctk.CTkFont(family="Arial", size=12, weight="bold")).pack(side="left")
        self.mode_menu = ctk.CTkOptionMenu(
            row_mode, values=list(MODES.keys()), variable=self.mode_var,
            width=200, font=ctk.CTkFont(family="Arial", size=12),
            command=self._on_mode_change
        )
        self.mode_menu.pack(side="left", padx=5)

        # Path row
        row_out = ctk.CTkFrame(config_frame, fg_color="transparent")
        row_out.pack(fill="x", padx=12, pady=(5, 10))
        self.lbl_out = ctk.CTkLabel(row_out, text="Đầu ra:", width=70, anchor="w", font=ctk.CTkFont(family="Arial", size=12, weight="bold"))
        self.lbl_out.pack(side="left")
        
        self.entry_out = ctk.CTkEntry(
            row_out, textvariable=self.out_path, placeholder_text="Chọn nơi lưu...",
            font=ctk.CTkFont(family="Arial", size=12)
        )
        self.entry_out.pack(side="left", fill="x", expand=True, padx=5)
        
        btn_browse_out = ctk.CTkButton(
            row_out, text="Browse", width=65, height=28,
            font=ctk.CTkFont(family="Arial", size=11, weight="bold"), fg_color="#2980b9", hover_color="#3498db",
            command=self._browse_output
        )
        btn_browse_out.pack(side="right", padx=2)

        # Output Preview & Logs Workspace
        self.preview_box = ctk.CTkTextbox(
            right_pane, fg_color="#111115", text_color="#a0a0b2",
            font=ctk.CTkFont(family="Consolas", size=12),
            border_width=1, border_color="#2c2c35", corner_radius=8
        )
        self.preview_box.grid(row=2, column=0, sticky="nsew", padx=15, pady=8)
        self._write_preview("💡 HỆ THỐNG ĐÃ SẴN SÀNG\n\n- Kéo thả file Markdown, Excel, hoặc Word của bạn vào khung bên trái.\n- Trình trích xuất thông minh sẽ tự động chuyển đổi file thành văn bản Markdown để bạn có thể xem trực quan, sửa hoặc xóa ký tự tùy ý trước khi xuất sang định dạng đầu ra mới.")

        # Large Action Button & Status Info
        action_frame = ctk.CTkFrame(right_pane, fg_color="transparent")
        action_frame.grid(row=3, column=0, sticky="ew", padx=15, pady=(5, 12))
        action_frame.columnconfigure(0, weight=3)
        action_frame.columnconfigure(1, weight=2)

        self.btn_convert = ctk.CTkButton(
            action_frame, text="⚡ CHUYỂN ĐỔI & LƯU ⚡", height=48,
            font=ctk.CTkFont(family="Arial", size=13, weight="bold"),
            fg_color="#2ecc71", hover_color="#27ae60", text_color="#ffffff",
            command=self._run_conversion
        )
        self.btn_convert.grid(row=0, column=0, sticky="ew", padx=(0, 5), pady=(0, 5))

        self.btn_open_file = ctk.CTkButton(
            action_frame, text="📂 MỞ FILE VỪA TẠO", height=48,
            font=ctk.CTkFont(family="Arial", size=13, weight="bold"),
            fg_color="#1c1c24", hover_color="#2c2c35", text_color="#8a8a9e",
            state="disabled",
            command=self._open_generated_file
        )
        self.btn_open_file.grid(row=0, column=1, sticky="ew", padx=(5, 0), pady=(0, 5))

        self.status_lbl = ctk.CTkLabel(
            action_frame, text="Sẵn sàng xử lý", font=ctk.CTkFont(family="Arial", size=12, slant="italic"), text_color="gray"
        )
        self.status_lbl.grid(row=1, column=0, columnspan=2, sticky="ew")

    # ── Internal Actions & Event Handlers ─────────────────────────────────────

    def _cfg(self):
        return MODES[self.mode_var.get()]

    def _on_mode_change(self, _=None):
        cfg = self._cfg()
        self.lbl_out.configure(text=cfg["out_label"] + ":")
        
        # Proactively update the output file path if input exists
        inp = self.in_path.get().strip()
        if inp:
            self._auto_output(inp)
        else:
            self.out_path.set("")
            
        self._set_status("Chế độ chuyển sang: " + self.mode_var.get(), "gray")

    def _clear_editor(self):
        self.editor.delete("1.0", "end")
        self.in_path.set("")
        self._update_counts()
        self._write_preview("📝 Trình soạn thảo đã trống. Bạn có thể tự viết Markdown tại đây!")
        self._set_status("Đã làm sạch không gian soạn thảo", "gray")
        self.drop_lbl.configure(text="📂 Kéo thả file hoặc nhấn 'Browse' để nạp nội dung...", text_color="#7eb8f5")

    def _update_counts(self, _=None):
        content = self.editor.get("1.0", "end-1c")
        chars = len(content)
        words = len(content.split())
        self.char_lbl.configure(text=f"Ký tự: {chars}")
        self.word_lbl.configure(text=f" |  Từ: {words}")

    def _auto_output(self, in_path: str):
        base = os.path.splitext(in_path)[0]
        self.out_path.set(base + self._cfg()["out_ext"])

    def _write_preview(self, msg: str):
        self.preview_box.configure(state="normal")
        self.preview_box.delete("1.0", "end")
        self.preview_box.insert("1.0", msg)
        self.preview_box.configure(state="disabled")

    def _set_status(self, msg: str, color: str = "white"):
        hex_colors = {"green": "#2ecc71", "red": "#e74c3c", "gray": "#8a8a9e", "orange": "#e67e22"}
        target_color = hex_colors.get(color, color)
        self.status_lbl.configure(text=msg, text_color=target_color)

    # ── Load file on drop or browse ───────────────────────────────────────────

    def _on_drop(self, event):
        path = event.data.strip().strip("{}")
        self._load_file_to_editor(path)

    def _browse_input(self):
        # We allow loading any supported document type to auto-detect
        path = filedialog.askopenfilename(filetypes=[
            ("Tài liệu được hỗ trợ", "*.md *.xlsx *.xls *.docx"),
            ("Markdown (*.md)", "*.md"),
            ("Excel (*.xlsx, *.xls)", "*.xlsx *.xls"),
            ("Word (*.docx)", "*.docx"),
            ("All Files", "*.*")
        ])
        if path:
            self._load_file_to_editor(path)

    def _load_file_to_editor(self, path: str):
        if not os.path.exists(path):
            self._set_status("❌ File nạp vào không tồn tại!", "red")
            return

        self._set_status("⏳ Đang nạp và trích xuất nội dung file...", "orange")
        ext = os.path.splitext(path)[1].lower()

        # Update paths
        self.in_path.set(path)
        self.drop_lbl.configure(text=f"✅ Nạp thành công: {os.path.basename(path)}", text_color="#2ecc71")

        def task():
            try:
                content = ""
                detected_mode = None

                if ext == ".md":
                    detected_mode = "MD → Excel"  # Default
                    with open(path, encoding="utf-8") as f:
                        content = f.read()
                elif ext in (".xlsx", ".xls"):
                    detected_mode = "Excel → MD"
                    content = extract_excel_to_md(path)
                elif ext == ".docx":
                    detected_mode = "Word → MD"
                    content = extract_word_to_md(path)
                else:
                    raise ValueError(f"Định dạng đuôi file {ext} không được hỗ trợ!")

                # Update UI on main thread
                def update_ui():
                    if detected_mode:
                        self.mode_var.set(detected_mode)
                        self._on_mode_change()
                    
                    self.editor.delete("1.0", "end")
                    self.editor.insert("1.0", content)
                    self._update_counts()
                    
                    # Generate Overview log
                    preview_msg = f"📂 TỔNG QUAN TÀI LIỆU ĐÃ NẠP:\n" \
                                  f"-----------------------------------\n" \
                                  f"- Tên file: {os.path.basename(path)}\n" \
                                  f"- Định dạng gốc: {ext.upper()}\n" \
                                  f"- Dung lượng: {os.path.getsize(path)} bytes\n" \
                                  f"- Tự động chuyển chế độ: {detected_mode}\n\n" \
                                  f"👉 Bây giờ bạn có thể đọc, sửa, thêm hoặc xóa ký tự tùy ý ở cột soạn thảo bên trái. Bấm 'Browse' ở mục Đầu ra nếu muốn đổi thư mục lưu mới."
                    
                    self._write_preview(preview_msg)
                    self._set_status("Nạp và trích xuất tài liệu thành công!", "green")

                self.after(0, update_ui)
            except Exception as e:
                self.after(0, lambda: self._set_status(f"❌ Lỗi nạp file: {e}", "red"))
                self.after(0, lambda: self._write_preview(f"❌ Chi tiết lỗi nạp:\n{e}"))

        threading.Thread(target=task, daemon=True).start()

    def _browse_output(self):
        ext = self._cfg()["out_ext"]
        path = filedialog.asksaveasfilename(
            defaultextension=ext, filetypes=OUT_FILETYPES[ext],
            initialfile=os.path.splitext(os.path.basename(self.in_path.get() or "document"))[0] + ext
        )
        if path:
            self.out_path.set(path)

    # ── Execute Document Generation & Export ─────────────────────────────────

    def _run_conversion(self):
        content = self.editor.get("1.0", "end-1c").strip()
        out = self.out_path.get().strip()
        mode = self.mode_var.get()

        if not content:
            self._set_status("⚠️ Trình soạn thảo trống, không có nội dung để chuyển!", "orange")
            return
        if not out:
            self._set_status("⚠️ Vui lòng chọn đường dẫn lưu tệp đầu ra!", "orange")
            return

        self.btn_convert.configure(state="disabled", text="⚡ Đang chuyển đổi...")
        self.btn_open_file.configure(state="disabled", fg_color="#1c1c24", text_color="#8a8a9e")
        self._set_status("⏳ Đang xử lý chuyển đổi và ghi file...", "orange")

        def task():
            try:
                # Direct dispatcher
                if mode == "MD → Excel":
                    msg = md_to_excel_from_text(content, out)
                elif mode == "MD → Word":
                    msg = md_to_word_from_text(content, out)
                elif mode in ("Excel → MD", "Word → MD"):
                    msg = save_markdown_from_text(content, out)
                else:
                    raise ValueError(f"Chế độ {mode} không hợp lệ!")

                color = "green" if msg.startswith("✅") else "red"
                
                # Output Overview details
                log_details = f"📈 NHẬT KÝ XUẤT BẢN TÀI LIỆU:\n" \
                              f"-----------------------------------\n" \
                              f"- Kết quả: THÀNH CÔNG 🎉\n" \
                              f"- Chế độ chuyển đổi: {mode}\n" \
                              f"- Tệp tin tạo ra: {os.path.basename(out)}\n" \
                              f"- Thư mục lưu: {os.path.dirname(out)}\n" \
                              f"- Dung lượng mới: {os.path.getsize(out) if os.path.exists(out) else 0} bytes\n\n" \
                              f"👉 Bấm nút '📂 MỞ FILE VỪA TẠO' bên dưới để mở và xem nội dung file trực tiếp!"

                self.after(0, lambda: self._write_preview(log_details))
                self.after(0, lambda: self._set_status(msg, color))
                
                if msg.startswith("✅"):
                    self.after(0, lambda: self.btn_open_file.configure(
                        state="normal", fg_color="#3a86ff", hover_color="#2563eb", text_color="#ffffff"
                    ))
            except Exception as e:
                self.after(0, lambda: self._set_status(f"❌ Lỗi chuyển đổi: {e}", "red"))
                self.after(0, lambda: self._write_preview(f"❌ Chi tiết lỗi chuyển đổi:\n{e}"))
            
            self.after(0, lambda: self.btn_convert.configure(state="normal", text="⚡ CHUYỂN ĐỔI & LƯU ⚡"))

        threading.Thread(target=task, daemon=True).start()

    def _open_generated_file(self):
        out = self.out_path.get().strip()
        if out and os.path.exists(out):
            try:
                os.startfile(out)
                self._set_status("🚀 Đã mở tệp tin bằng ứng dụng mặc định", "green")
            except Exception as e:
                self._set_status(f"❌ Không thể mở file: {e}", "red")
        else:
            self._set_status("⚠️ Tệp tin chưa được tạo hoặc không tồn tại!", "orange")


if __name__ == "__main__":
    app = App()
    app.mainloop()